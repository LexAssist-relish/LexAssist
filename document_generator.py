import os
import json
import logging
import base64
from typing import Dict, Any, Optional, List
from datetime import datetime
from io import BytesIO
from flask import send_file

# PDF generation libraries
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('document_generator')

class DocumentGenerator:
    """
    Generates downloadable documents from legal analysis results.
    """
    
    def __init__(self, logo_path: Optional[str] = None):
        """
        Initialize the document generator.
        
        Args:
            logo_path: Path to the logo image file
        """
        self.logo_path = logo_path or '/home/ubuntu/legal_app_frontend/public/images/logo.png'
        logger.info("Document generator initialized")
    
    def generate_pdf(self, brief_text: str, analysis_results: Dict[str, Any], 
                   output_path: Optional[str] = None) -> str:
        """
        Generate a PDF document from the analysis results.
        
        Args:
            brief_text: The original brief text
            analysis_results: The analysis results
            output_path: Path to save the PDF file (optional)
            
        Returns:
            Path to the generated PDF file
        """
        logger.info("Generating PDF document")
        
        # Create a file-like object to receive PDF data
        buffer = BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            buffer if output_path is None else output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = styles['Title']
        heading_style = styles['Heading1']
        subheading_style = styles['Heading2']
        normal_style = styles['Normal']
        
        # Create custom styles
        section_style = ParagraphStyle(
            'SectionStyle',
            parent=styles['Heading3'],
            textColor=colors.navy,
            spaceAfter=6
        )
        
        # Create story (content)
        story = []
        
        # Add logo if available
        if os.path.exists(self.logo_path):
            img = Image(self.logo_path, width=2*72, height=1*72)  # 2x1 inches
            img.hAlign = 'CENTER'
            story.append(img)
            story.append(Spacer(1, 12))
        
        # Add title
        story.append(Paragraph("Legal Analysis Report", title_style))
        story.append(Spacer(1, 12))
        
        # Add date
        current_date = datetime.now().strftime("%d %B %Y")
        story.append(Paragraph(f"Generated on: {current_date}", normal_style))
        story.append(Spacer(1, 24))
        
        # Add brief section
        story.append(Paragraph("Case Brief", heading_style))
        story.append(Paragraph(brief_text, normal_style))
        story.append(Spacer(1, 24))
        
        # Add summary section if available
        if 'analysis' in analysis_results and 'summary' in analysis_results['analysis']:
            story.append(Paragraph("Summary", heading_style))
            story.append(Paragraph(analysis_results['analysis']['summary'], normal_style))
            story.append(Spacer(1, 24))
        
        # Add law sections if available
        if 'lawSections' in analysis_results and analysis_results['lawSections']:
            story.append(Paragraph("Relevant Law Sections", heading_style))
            
            for section in analysis_results['lawSections']:
                title = section.get('title', 'Unknown Act')
                section_number = section.get('sectionNumber', 'N/A')
                content = section.get('content', '')
                
                story.append(Paragraph(f"{title}, Section {section_number}", section_style))
                story.append(Paragraph(content, normal_style))
                story.append(Spacer(1, 12))
            
            story.append(Spacer(1, 12))
        
        # Add case histories if available
        if 'caseHistories' in analysis_results and analysis_results['caseHistories']:
            story.append(Paragraph("Relevant Case Histories", heading_style))
            
            for case in analysis_results['caseHistories']:
                citation = case.get('citation', 'Unknown Citation')
                parties = case.get('parties', 'Unknown Parties')
                holdings = case.get('holdings', '')
                date = case.get('date', '')
                
                story.append(Paragraph(f"{parties} ({citation})", section_style))
                if date:
                    story.append(Paragraph(f"Date: {date}", normal_style))
                story.append(Paragraph(f"Holdings: {holdings}", normal_style))
                story.append(Spacer(1, 12))
            
            story.append(Spacer(1, 12))
        
        # Add arguments if available
        if 'analysis' in analysis_results and 'arguments' in analysis_results['analysis']:
            story.append(Paragraph("Legal Arguments", heading_style))
            
            arguments = analysis_results['analysis']['arguments']
            for i, argument in enumerate(arguments, 1):
                story.append(Paragraph(f"{i}. {argument}", normal_style))
                story.append(Spacer(1, 6))
            
            story.append(Spacer(1, 12))
        
        # Add challenges if available
        if 'analysis' in analysis_results and 'challenges' in analysis_results['analysis']:
            story.append(Paragraph("Potential Challenges", heading_style))
            
            challenges = analysis_results['analysis']['challenges']
            for i, challenge in enumerate(challenges, 1):
                story.append(Paragraph(f"{i}. {challenge}", normal_style))
                story.append(Spacer(1, 6))
            
            story.append(Spacer(1, 12))
        
        # Add recommendations if available
        if 'analysis' in analysis_results and 'recommendations' in analysis_results['analysis']:
            story.append(Paragraph("Recommendations", heading_style))
            
            recommendations = analysis_results['analysis']['recommendations']
            for i, recommendation in enumerate(recommendations, 1):
                story.append(Paragraph(f"{i}. {recommendation}", normal_style))
                story.append(Spacer(1, 6))
        
        # Add disclaimer
        story.append(Spacer(1, 36))
        disclaimer_style = ParagraphStyle(
            'Disclaimer',
            parent=normal_style,
            fontSize=8,
            textColor=colors.gray
        )
        disclaimer_text = (
            "DISCLAIMER: This analysis is generated by Lex Assist AI and is provided for informational purposes only. "
            "It does not constitute legal advice and should not be relied upon as such. "
            "Please consult with a qualified legal professional for advice specific to your situation."
        )
        story.append(Paragraph(disclaimer_text, disclaimer_style))
        
        # Build the PDF
        doc.build(story)
        
        # If output_path is provided, return it
        if output_path:
            return output_path
        
        # Otherwise, get the PDF data from the buffer and save it to a temporary file
        buffer.seek(0)
        temp_path = f"/tmp/legal_analysis_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
        with open(temp_path, 'wb') as f:
            f.write(buffer.read())
        
        return temp_path
    
    def generate_docx(self, brief_text: str, analysis_results: Dict[str, Any], 
                    output_path: Optional[str] = None) -> str:
        """
        Generate a DOCX document from the analysis results.
        
        Args:
            brief_text: The original brief text
            analysis_results: The analysis results
            output_path: Path to save the DOCX file (optional)
            
        Returns:
            Path to the generated DOCX file
        """
        logger.info("Generating DOCX document")
        
        # Create a new Document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Legal Analysis Report"
        doc.core_properties.author = "Lex Assist AI"
        
        # Add title
        title = doc.add_heading("Legal Analysis Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime("%d %B %Y")
        date_paragraph = doc.add_paragraph(f"Generated on: {current_date}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add space
        
        # Add brief section
        doc.add_heading("Case Brief", level=1)
        doc.add_paragraph(brief_text)
        
        doc.add_paragraph()  # Add space
        
        # Add summary section if available
        if 'analysis' in analysis_results and 'summary' in analysis_results['analysis']:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(analysis_results['analysis']['summary'])
            doc.add_paragraph()  # Add space
        
        # Add law sections if available
        if 'lawSections' in analysis_results and analysis_results['lawSections']:
            doc.add_heading("Relevant Law Sections", level=1)
            
            for section in analysis_results['lawSections']:
                title = section.get('title', 'Unknown Act')
                section_number = section.get('sectionNumber', 'N/A')
                content = section.get('content', '')
                
                heading = doc.add_heading(f"{title}, Section {section_number}", level=2)
                heading.style.font.color.rgb = RGBColor(0, 32, 96)  # Navy blue
                doc.add_paragraph(content)
            
            doc.add_paragraph()  # Add space
        
        # Add case histories if available
        if 'caseHistories' in analysis_results and analysis_results['caseHistories']:
            doc.add_heading("Relevant Case Histories", level=1)
            
            for case in analysis_results['caseHistories']:
                citation = case.get('citation', 'Unknown Citation')
                parties = case.get('parties', 'Unknown Parties')
                holdings = case.get('holdings', '')
                date = case.get('date', '')
                
                heading = doc.add_heading(f"{parties} ({citation})", level=2)
                heading.style.font.color.rgb = RGBColor(0, 32, 96)  # Navy blue
                
                if date:
                    date_para = doc.add_paragraph()
                    date_para.add_run(f"Date: ").bold = True
                    date_para.add_run(date)
                
                holdings_para = doc.add_paragraph()
                holdings_para.add_run(f"Holdings: ").bold = True
                holdings_para.add_run(holdings)
            
            doc.add_paragraph()  # Add space
        
        # Add arguments if available
        if 'analysis' in analysis_results and 'arguments' in analysis_results['analysis']:
            doc.add_heading("Legal Arguments", level=1)
            
            arguments = analysis_results['analysis']['arguments']
            for i, argument in enumerate(arguments, 1):
                doc.add_paragraph(f"{i}. {argument}", style='List Number')
            
            doc.add_paragraph()  # Add space
        
        # Add challenges if available
        if 'analysis' in analysis_results and 'challenges' in analysis_results['analysis']:
            doc.add_heading("Potential Challenges", level=1)
            
            challenges = analysis_results['analysis']['challenges']
            for i, challenge in enumerate(challenges, 1):
                doc.add_paragraph(f"{i}. {challenge}", style='List Number')
            
            doc.add_paragraph()  # Add space
        
        # Add recommendations if available
        if 'analysis' in analysis_results and 'recommendations' in analysis_results['analysis']:
            doc.add_heading("Recommendations", level=1)
            
            recommendations = analysis_results['analysis']['recommendations']
            for i, recommendation in enumerate(recommendations, 1):
                doc.add_paragraph(f"{i}. {recommendation}", style='List Number')
        
        # Add disclaimer
        doc.add_paragraph()  # Add space
        disclaimer = doc.add_paragraph(
            "DISCLAIMER: This analysis is generated by Lex Assist AI and is provided for informational purposes only. "
            "It does not constitute legal advice and should not be relied upon as such. "
            "Please consult with a qualified legal professional for advice specific to your situation."
        )
        disclaimer.style = doc.styles['Normal']
        disclaimer.style.font.size = Pt(8)
        disclaimer.style.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        
        # Determine output path
        if output_path is None:
            output_path = f"/tmp/legal_analysis_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        
        # Save the document
        doc.save(output_path)
        
        return output_path
    
    def generate_txt(self, brief_text: str, analysis_results: Dict[str, Any], 
                   output_path: Optional[str] = None) -> str:
        """
        Generate a TXT document from the analysis results.
        
        Args:
            brief_text: The original brief text
            analysis_results: The analysis results
            output_path: Path to save the TXT file (optional)
            
        Returns:
            Path to the generated TXT file
        """
        logger.info("Generating TXT document")
        
        # Create content
        content = []
        
        # Add title
        content.append("LEGAL ANALYSIS REPORT")
        content.append("=" * 80)
        content.append("")
        
        # Add date
        current_date = datetime.now().strftime("%d %B %Y")
        content.append(f"Generated on: {current_date}")
        content.append("")
        content.append("-" * 80)
        content.append("")
        
        # Add brief section
        content.append("CASE BRIEF")
        content.append("-" * 80)
        content.append(brief_text)
        content.append("")
        content.append("-" * 80)
        content.append("")
        
        # Add summary section if available
        if 'analysis' in analysis_results and 'summary' in analysis_results['analysis']:
            content.append("SUMMARY")
            content.append("-" * 80)
            content.append(analysis_results['analysis']['summary'])
            content.append("")
            content.append("-" * 80)
            content.append("")
        
        # Add law sections if available
        if 'lawSections' in analysis_results and analysis_results['lawSections']:
            content.append("RELEVANT LAW SECTIONS")
            content.append("-" * 80)
            
            for section in analysis_results['lawSections']:
                title = section.get('title', 'Unknown Act')
                section_number = section.get('sectionNumber', 'N/A')
                section_content = section.get('content', '')
                
                content.append(f"{title}, Section {section_number}")
                content.append(section_content)
                content.append("")
            
            content.append("-" * 80)
            content.append("")
        
        # Add case histories if available
        if 'caseHistories' in analysis_results and analysis_results['caseHistories']:
            content.append("RELEVANT CASE HISTORIES")
            content.append("-" * 80)
            
            for case in analysis_results['caseHistories']:
                citation = case.get('citation', 'Unknown Citation')
                parties = case.get('parties', 'Unknown Parties')
                holdings = case.get('holdings', '')
                date = case.get('date', '')
                
                content.append(f"{parties} ({citation})")
                if date:
                    content.append(f"Date: {date}")
                content.append(f"Holdings: {holdings}")
                content.append("")
            
            content.append("-" * 80)
            content.append("")
        
        # Add arguments if available
        if 'analysis' in analysis_results and 'arguments' in analysis_results['analysis']:
            content.append("LEGAL ARGUMENTS")
            content.append("-" * 80)
            
            arguments = analysis_results['analysis']['arguments']
            for i, argument in enumerate(arguments, 1):
                content.append(f"{i}. {argument}")
            
            content.append("")
            content.append("-" * 80)
            content.append("")
        
        # Add challenges if available
        if 'analysis' in analysis_results and 'challenges' in analysis_results['analysis']:
            content.append("POTENTIAL CHALLENGES")
            content.append("-" * 80)
            
            challenges = analysis_results['analysis']['challenges']
            for i, challenge in enumerate(challenges, 1):
                content.append(f"{i}. {challenge}")
            
            content.append("")
            content.append("-" * 80)
            content.append("")
        
        # Add recommendations if available
        if 'analysis' in analysis_results and 'recommendations' in analysis_results['analysis']:
            content.append("RECOMMENDATIONS")
            content.append("-" * 80)
            
            recommendations = analysis_results['analysis']['recommendations']
            for i, recommendation in enumerate(recommendations, 1):
                content.append(f"{i}. {recommendation}")
            
            content.append("")
            content.append("-" * 80)
            content.append("")
        
        # Add disclaimer
        content.append("DISCLAIMER: This analysis is generated by Lex Assist AI and is provided for")
        content.append("informational purposes only. It does not constitute legal advice and should")
        content.append("not be relied upon as such. Please consult with a qualified legal professional")
        content.append("for advice specific to your situation.")
        
        # Join content with newlines
        full_content = "\n".join(content)
        
        # Determine output path
        if output_path is None:
            output_path = f"/tmp/legal_analysis_{datetime.now().strftime('%Y%m%d%H%M%S')}.txt"
        
        # Write to file
        with open(output_path, 'w') as f:
            f.write(full_content)
        
        return output_path
    
    def generate_email_content(self, brief_text: str, analysis_results: Dict[str, Any]) -> Dict[str, str]:
        """
        Generate email content for sharing analysis results.
        
        Args:
            brief_text: The original brief text
            analysis_results: The analysis results
            
        Returns:
            Dict containing email subject and body
        """
        logger.info("Generating email content")
        
        # Create email subject
        subject = "Legal Analysis Report from Lex Assist"
        
        # Create email body
        body = []
        body.append("Please find below a legal analysis report generated by Lex Assist AI:")
        body.append("")
        
        # Add summary if available
        if 'analysis' in analysis_results and 'summary' in analysis_results['analysis']:
            body.append("SUMMARY:")
            body.append(analysis_results['analysis']['summary'])
            body.append("")
        
        # Add key law sections if available
        if 'lawSections' in analysis_results and analysis_results['lawSections']:
            body.append("KEY LEGAL PROVISIONS:")
            
            for section in analysis_results['lawSections'][:2]:  # Limit to top 2
                title = section.get('title', 'Unknown Act')
                section_number = section.get('sectionNumber', 'N/A')
                
                body.append(f"- {title}, Section {section_number}")
            
            body.append("")
        
        # Add key case histories if available
        if 'caseHistories' in analysis_results and analysis_results['caseHistories']:
            body.append("KEY CASE PRECEDENTS:")
            
            for case in analysis_results['caseHistories'][:2]:  # Limit to top 2
                citation = case.get('citation', 'Unknown Citation')
                parties = case.get('parties', 'Unknown Parties')
                
                body.append(f"- {parties} ({citation})")
            
            body.append("")
        
        # Add note about attachment
        body.append("A complete analysis report is attached to this email.")
        body.append("")
        
        # Add disclaimer
        body.append("DISCLAIMER: This analysis is generated by Lex Assist AI and is provided for informational purposes only.")
        body.append("It does not constitute legal advice and should not be relied upon as such.")
        body.append("Please consult with a qualified legal professional for advice specific to your situation.")
        
        # Join body with newlines
        full_body = "\n".join(body)
        
        return {
            "subject": subject,
            "body": full_body
        }
    
    def generate_whatsapp_content(self, brief_text: str, analysis_results: Dict[str, Any]) -> str:
        """
        Generate WhatsApp content for sharing analysis results.
        
        Args:
            brief_text: The original brief text
            analysis_results: The analysis results
            
        Returns:
            WhatsApp message content
        """
        logger.info("Generating WhatsApp content")
        
        # Create WhatsApp message
        message = []
        message.append("*Legal Analysis Report from Lex Assist*")
        message.append("")
        
        # Add summary if available
        if 'analysis' in analysis_results and 'summary' in analysis_results['analysis']:
            message.append("*Summary:*")
            message.append(analysis_results['analysis']['summary'])
            message.append("")
        
        # Add key recommendations if available
        if 'analysis' in analysis_results and 'recommendations' in analysis_results['analysis']:
            message.append("*Key Recommendations:*")
            
            recommendations = analysis_results['analysis']['recommendations']
            for i, recommendation in enumerate(recommendations[:3], 1):  # Limit to top 3
                message.append(f"{i}. {recommendation}")
            
            message.append("")
        
        # Add note about full report
        message.append("_A complete analysis report has been generated and is available in the Lex Assist app._")
        message.append("")
        
        # Add disclaimer
        message.append("_DISCLAIMER: This analysis is provided for informational purposes only and does not constitute legal advice._")
        
        # Join message with newlines
        full_message = "\n".join(message)
        
        return full_message


# Example usage
if __name__ == "__main__":
    generator = DocumentGenerator()
    
    # Example brief and analysis results
    brief = "This case involves a breach of contract under Section 73 of the Indian Contract Act."
    analysis = {
        "lawSections": [
            {
                "title": "Indian Contract Act",
                "sectionNumber": "73",
                "content": "When a contract has been broken, the party who suffers by such breach is entitled to receive compensation."
            }
        ],
        "caseHistories": [
            {
                "citation": "AIR 2017 SC 567",
                "parties": "Mehta vs. Patel & Others",
                "holdings": "The Court established that in cases of contractual breach, the aggrieved party must prove actual damages.",
                "date": "05 Aug 2017"
            }
        ],
        "analysis": {
            "summary": "This case involves a breach of contract under Section 73 of the Indian Contract Act.",
            "arguments": ["The defendant's actions constitute a clear breach of contract."],
            "challenges": ["The defense may argue force majeure."],
            "recommendations": ["Gather all communication records to establish the terms of the contract."]
        }
    }
    
    # Generate documents
    pdf_path = generator.generate_pdf(brief, analysis)
    docx_path = generator.generate_docx(brief, analysis)
    txt_path = generator.generate_txt(brief, analysis)
    
    print(f"PDF generated: {pdf_path}")
    print(f"DOCX generated: {docx_path}")
    print(f"TXT generated: {txt_path}")
    
    # Generate sharing content
    email_content = generator.generate_email_content(brief, analysis)
    whatsapp_content = generator.generate_whatsapp_content(brief, analysis)
    
    print("\nEmail Subject:", email_content["subject"])
    print("Email Body:", email_content["body"])
    print("\nWhatsApp Message:", whatsapp_content)
